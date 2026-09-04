#!/usr/bin/env python
"""Build résumé variants from Markdown sources into every format we hand out.

Source of truth: one Markdown file per variant in ``src/``. The Markdown is
deliberately strict (see README.md) so it stays readable by a person, easy to
edit, and trivially parseable into a data model. From that model we render:

    md        normalised copy of the source (canonical, human + machine)
    json      JSON Resume schema (https://jsonresume.org/schema) for ATS/agents
    txt       plain text, no markup, for ATS upload boxes and LLM ingestion
    html      standalone page with theme.css inlined
    pdf       printed from the HTML through Chromium (text stays selectable)
    jpg       full-page screenshot of the HTML for previews and image uploads
    docx      Word document built with python-docx, for job portals
    linkedin  paste-ready blocks sized to LinkedIn's field limits

Fail-closed guard: every term in ``denylist.txt`` is checked against the source
and every text output before anything is written. A hit aborts the build.

Usage:
    python build.py                     # build every variant, every format
    python build.py general cloud       # only these variants
    python build.py --formats pdf,json  # only these formats
    python build.py --check             # denylist scan only, no output
    python build.py --publish general   # also copy that PDF to dist/ for the site
"""

from __future__ import annotations

import argparse
import html
import io
import json
import re
import shutil
import sys
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

HERE = Path(__file__).resolve().parent
SRC = HERE / "src"
OUT = HERE / "out"
DIST = HERE / "dist"
THEME = HERE / "theme.css"
DENYLIST = HERE / "denylist.txt"
PUBLISHED_NAME = "Steven-Shine-Resume.pdf"

ALL_FORMATS = ["md", "json", "txt", "html", "pdf", "jpg", "docx", "linkedin"]


# --------------------------------------------------------------------------- #
# Data model
# --------------------------------------------------------------------------- #

@dataclass
class Entry:
    title: str
    org: str = ""
    dates: str = ""
    location: str = ""
    description: list[str] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)


@dataclass
class ListItem:
    label: str | None
    text: str


@dataclass
class Section:
    title: str
    kind: str                       # "text" | "list" | "entries"
    paragraphs: list[str] = field(default_factory=list)
    items: list[ListItem] = field(default_factory=list)
    entries: list[Entry] = field(default_factory=list)


@dataclass
class Resume:
    meta: dict[str, str]
    sections: list[Section]

    @property
    def name(self) -> str:
        return self.meta.get("name", "")

    def section(self, title: str) -> Section | None:
        for s in self.sections:
            if s.title.lower() == title.lower():
                return s
        return None


# --------------------------------------------------------------------------- #
# Parser
# --------------------------------------------------------------------------- #

_FRONT = re.compile(r"^---\s*$")
_KV = re.compile(r"^([A-Za-z_][\w-]*):\s*(.*)$")
_BULLET = re.compile(r"^- (.*)$")
_LABELLED = re.compile(r"^\*\*(.+?):\*\*\s*(.*)$")


def parse(text: str, source: str = "<memory>") -> Resume:
    lines = text.replace("\r\n", "\n").split("\n")
    i = 0
    meta: dict[str, str] = {}

    if lines and _FRONT.match(lines[0]):
        i = 1
        while i < len(lines) and not _FRONT.match(lines[i]):
            m = _KV.match(lines[i])
            if m:
                meta[m.group(1).strip()] = m.group(2).strip()
            elif lines[i].strip():
                raise ValueError(f"{source}: bad front-matter line {i + 1}: {lines[i]!r}")
            i += 1
        i += 1
    else:
        raise ValueError(f"{source}: missing front matter (--- block) at top of file")

    sections: list[Section] = []
    cur: Section | None = None
    entry: Entry | None = None
    saw_meta_line = False

    def flush_bullet(target: list[str], buf: list[str]) -> None:
        if buf:
            target.append(" ".join(s.strip() for s in buf))
            buf.clear()

    bullet_buf: list[str] = []
    bullet_target: list[str] | None = None

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        i += 1

        if line.startswith("## "):
            flush_bullet(bullet_target if bullet_target is not None else [], bullet_buf)
            bullet_target = None
            cur = Section(title=line[3:].strip(), kind="text")
            sections.append(cur)
            entry = None
            continue

        if cur is None:
            if line.strip():
                raise ValueError(f"{source}: content before first '## ' section: {line!r}")
            continue

        if line.startswith("### "):
            flush_bullet(bullet_target if bullet_target is not None else [], bullet_buf)
            cur.kind = "entries"
            head = line[4:].strip()
            title, _, org = head.partition(" | ")
            entry = Entry(title=title.strip(), org=org.strip())
            cur.entries.append(entry)
            saw_meta_line = False
            bullet_target = None
            continue

        m = _BULLET.match(line)
        if m:
            flush_bullet(bullet_target if bullet_target is not None else [], bullet_buf)
            if entry is not None:
                bullet_target = entry.bullets
                bullet_buf = [m.group(1)]
            else:
                cur.kind = "list"
                lm = _LABELLED.match(m.group(1))
                cur.items.append(ListItem(lm.group(1), lm.group(2)) if lm else ListItem(None, m.group(1)))
                bullet_target = None
            continue

        if raw.startswith("  ") and bullet_target is not None and bullet_buf:
            bullet_buf.append(line)                # bullet continuation
            continue

        flush_bullet(bullet_target if bullet_target is not None else [], bullet_buf)
        bullet_target = None

        if not line.strip():
            continue

        if entry is not None:
            if not saw_meta_line and not entry.bullets:
                dates, _, loc = line.partition(" · ")
                entry.dates, entry.location = dates.strip(), loc.strip()
                saw_meta_line = True
            else:
                entry.description.append(line.strip())
        else:
            cur.paragraphs.append(line.strip())

    flush_bullet(bullet_target if bullet_target is not None else [], bullet_buf)
    return Resume(meta=meta, sections=sections)


# --------------------------------------------------------------------------- #
# Inline formatting helpers
# --------------------------------------------------------------------------- #

_BOLD = re.compile(r"\*\*(.+?)\*\*")


def strip_inline(s: str) -> str:
    return _BOLD.sub(r"\1", s)


def inline_html(s: str) -> str:
    parts = _BOLD.split(s)
    out = []
    for n, p in enumerate(parts):
        p = html.escape(p)
        out.append(f"<strong>{p}</strong>" if n % 2 else p)
    return "".join(out)


def inline_runs(s: str) -> list[tuple[str, bool]]:
    parts = _BOLD.split(s)
    return [(p, bool(n % 2)) for n, p in enumerate(parts) if p]


# --------------------------------------------------------------------------- #
# Denylist
# --------------------------------------------------------------------------- #

def load_denylist() -> list[str]:
    if not DENYLIST.exists():
        return []
    terms = []
    for line in DENYLIST.read_text(encoding="utf-8").splitlines():
        line = line.split("#", 1)[0].strip()
        if line:
            terms.append(line)
    return terms


def denylist_hits(text: str, terms: list[str]) -> list[str]:
    low = text.lower()
    return [t for t in terms if t.lower() in low]


def guard(label: str, text: str, terms: list[str]) -> None:
    hits = denylist_hits(text, terms)
    if hits:
        sys.exit(f"DENYLIST HIT in {label}: {', '.join(hits)} — build aborted, nothing written.")


# --------------------------------------------------------------------------- #
# Renderers: text-ish
# --------------------------------------------------------------------------- #

def contact_items(r: Resume) -> list[tuple[str, str]]:
    keys = [("email", "mailto:{}"), ("phone", ""), ("website", "https://{}"),
            ("linkedin", "https://{}"), ("github", "https://{}")]
    items = []
    for k, href in keys:
        v = r.meta.get(k)
        if v:
            items.append((v, href.format(v) if href else ""))
    return items


def render_txt(r: Resume) -> str:
    o = [r.name.upper()]
    tl = " · ".join(x for x in [r.meta.get("title"), r.meta.get("location")] if x)
    if tl:
        o.append(tl)
    o.append(" | ".join(v for v, _ in contact_items(r)))
    for s in r.sections:
        o += ["", s.title.upper(), "-" * len(s.title)]
        if s.kind == "text":
            o += [strip_inline(p) for p in s.paragraphs]
        elif s.kind == "list":
            for it in s.items:
                o.append(f"{it.label}: {strip_inline(it.text)}" if it.label else strip_inline(it.text))
        else:
            for e in s.entries:
                head = e.title + (f" — {e.org}" if e.org else "")
                o.append(head)
                sub = " · ".join(x for x in [e.dates, e.location] if x)
                if sub:
                    o.append(sub)
                o += [strip_inline(d) for d in e.description]
                o += [f"  - {strip_inline(b)}" for b in e.bullets]
                o.append("")
    return "\n".join(o).rstrip() + "\n"


_MONTHS = {m.lower(): n for n, m in enumerate(
    ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"], 1)}


def iso_date(s: str) -> str | None:
    s = s.strip()
    if not s or s.lower() in ("present", "current", "now"):
        return None
    m = re.match(r"([A-Za-z]{3})[a-z]*\.?\s+(\d{4})$", s)
    if m and m.group(1).lower() in _MONTHS:
        return f"{m.group(2)}-{_MONTHS[m.group(1).lower()]:02d}"
    m = re.match(r"(\d{4})$", s)
    return m.group(1) if m else None


def split_dates(d: str) -> tuple[str | None, str | None]:
    parts = re.split(r"\s+[–—-]\s+", d, maxsplit=1)
    start = iso_date(parts[0]) if parts else None
    end = iso_date(parts[1]) if len(parts) > 1 else None
    return start, end


def render_json(r: Resume) -> str:
    city, _, region = r.meta.get("location", "").partition(", ")
    profiles = []
    for net, label in (("linkedin", "LinkedIn"), ("github", "GitHub")):
        v = r.meta.get(net)
        if v:
            profiles.append({"network": label, "username": v.rstrip("/").split("/")[-1],
                             "url": f"https://{v}"})
    basics = {
        "name": r.name,
        "label": r.meta.get("title", ""),
        "email": r.meta.get("email", ""),
        "phone": r.meta.get("phone", ""),
        "url": f"https://{r.meta['website']}" if r.meta.get("website") else "",
        "summary": " ".join(strip_inline(p) for p in (r.section("Summary").paragraphs if r.section("Summary") else [])),
        "location": {"city": city, "region": region, "countryCode": "US"},
        "profiles": profiles,
    }
    doc: dict = {
        "$schema": "https://raw.githubusercontent.com/jsonresume/resume-schema/v1.0.0/schema.json",
        "basics": basics,
        "work": [], "projects": [], "skills": [], "education": [], "certificates": [],
        "meta": {"variant": r.meta.get("variant", ""), "target": r.meta.get("target", ""),
                 "canonical": r.meta.get("website", ""), "lastModified": date.today().isoformat(),
                 "format": "JSON Resume v1.0.0"},
    }
    for s in r.sections:
        key = s.title.lower()
        if key == "experience":
            for e in s.entries:
                start, end = split_dates(e.dates)
                w = {"name": e.org, "position": e.title, "location": e.location,
                     "summary": " ".join(strip_inline(d) for d in e.description),
                     "highlights": [strip_inline(b) for b in e.bullets]}
                if start:
                    w["startDate"] = start
                if end:
                    w["endDate"] = end
                doc["work"].append(w)
        elif key == "projects":
            for e in s.entries:
                start, end = split_dates(e.dates)
                p = {"name": e.title, "url": f"https://{e.org}" if e.org and "." in e.org and " " not in e.org else "",
                     "description": " ".join(strip_inline(d) for d in e.description),
                     "highlights": [strip_inline(b) for b in e.bullets]}
                if start:
                    p["startDate"] = start
                if end:
                    p["endDate"] = end
                doc["projects"].append(p)
        elif key == "skills":
            for it in s.items:
                doc["skills"].append({"name": it.label or "General",
                                      "keywords": [k.strip() for k in strip_inline(it.text).split(",") if k.strip()]})
        elif key == "education":
            for e in s.entries:
                study, _, area = e.title.partition(", ")
                _, end = split_dates(e.dates) if " – " in e.dates else (None, iso_date(e.dates))
                ed = {"institution": e.org, "studyType": study, "area": area}
                if end:
                    ed["endDate"] = end
                doc["education"].append(ed)
        elif key == "certifications":
            for it in s.items:
                doc["certificates"].append({"name": strip_inline(it.text), "issuer": it.label or ""})
    return json.dumps(doc, indent=2, ensure_ascii=False) + "\n"


_LEAD_NUM = re.compile(r"^(~?\$?\d[\d,.]*[+%]?)\s+(.*)$")


def render_highlights(r: Resume) -> list[str]:
    """Optional `highlights:` front matter -> a strip of pipe-separated facts.
    Rendered as text (no images), so extraction sees e.g. "9 domain controllers rebuilt"."""
    hl = r.meta.get("highlights")
    if not hl:
        return []
    out = ['<ul class="highlights">']
    for h in (x.strip() for x in hl.split("|")):
        if not h:
            continue
        m = _LEAD_NUM.match(h)
        if m:
            out.append(f'<li><strong class="n">{html.escape(m.group(1))}</strong> '
                       f'<span class="t">{html.escape(m.group(2))}</span></li>')
        else:
            out.append(f'<li><span class="t">{html.escape(h)}</span></li>')
    out.append("</ul>")
    return out


def render_html(r: Resume, css: str) -> str:
    def esc(s: str) -> str:
        return html.escape(s)

    def slug(s: str) -> str:
        return re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")

    body = [f'<header class="head"><h1>{esc(r.name)}</h1>']
    # Plain "|" separators: middle dots and dashes are the characters ATS text
    # extractors most often mangle, and the contact line is what they key on.
    tl = " | ".join(x for x in [r.meta.get("title"), r.meta.get("location")] if x)
    if tl:
        body.append(f'<p class="tagline">{esc(tl)}</p>')
    links = []
    for v, href in contact_items(r):
        links.append(f'<a href="{esc(href)}">{esc(v)}</a>' if href else f"<span>{esc(v)}</span>")
    body.append('<p class="contact">' + '<span class="sep">|</span>'.join(links) + "</p></header>")

    highlights = render_highlights(r)
    if highlights and not r.section("Summary"):
        body += highlights

    for s in r.sections:
        body.append(f'<section class="sec sec-{slug(s.title)}"><h2>{esc(s.title)}</h2>')
        if s.kind == "text":
            body += [f"<p>{inline_html(p)}</p>" for p in s.paragraphs]
        elif s.kind == "list":
            body.append('<ul class="skills">')
            for it in s.items:
                # The trailing space lives inside the label so flex layouts, which
                # drop whitespace-only text nodes, still extract "Label: value".
                lab = f'<strong class="k">{esc(it.label)}: </strong>' if it.label else ""
                body.append(f'<li>{lab}<span class="v">{inline_html(it.text)}</span></li>')
            body.append("</ul>")
        else:
            for n, e in enumerate(s.entries):
                body.append(f'<div class="entry{" first" if n == 0 else ""}">')
                body.append('<div class="row"><span class="title">' + esc(e.title)
                            + (f' <span class="org">— {esc(e.org)}</span>' if e.org else "")
                            + f'</span><span class="dates">{esc(e.dates)}</span></div>')
                if e.location:
                    body.append(f'<div class="loc">{esc(e.location)}</div>')
                body += [f"<p>{inline_html(d)}</p>" for d in e.description]
                if e.bullets:
                    body.append("<ul>" + "".join(f"<li>{inline_html(b)}</li>" for b in e.bullets) + "</ul>")
                body.append("</div>")
        body.append("</section>")
        if highlights and s.title.lower() == "summary":
            body += highlights

    desc = esc(r.meta.get("title", "")) + " résumé"
    return (
        "<!doctype html><html lang=\"en\"><head><meta charset=\"utf-8\">"
        f"<title>{esc(r.name)} — Résumé</title>"
        f'<meta name="author" content="{esc(r.name)}"><meta name="description" content="{desc}">'
        f"<style>{css}</style></head><body><main class=\"page\">" + "\n".join(body) + "</main></body></html>\n"
    )


LINKEDIN_LIMITS = {"headline": 220, "about": 2600, "description": 2000}


def linkedin_headline(r: Resume) -> str:
    return r.meta.get("headline") or r.meta.get("title", "")


def linkedin_about(r: Resume) -> str:
    summ = r.section("Summary")
    return "\n\n".join(strip_inline(p) for p in summ.paragraphs) if summ else ""


def linkedin_entry_text(e: Entry, limit: int | None = LINKEDIN_LIMITS["description"]) -> tuple[str, int]:
    """Description paragraphs, then '• ' bullets; bullets that would breach the limit are dropped.

    Returns (text, number of bullets dropped)."""
    lines = [strip_inline(d) for d in e.description]
    dropped = 0
    for b in e.bullets:
        candidate = "\n".join([*lines, f"• {strip_inline(b)}"])
        if limit is not None and len(candidate) > limit:
            dropped += 1
            continue
        lines.append(f"• {strip_inline(b)}")
    return "\n".join(lines), dropped


_PAREN = re.compile(r"\s*\([^)]*\)")


def split_top_level(text: str) -> list[str]:
    """Split on commas that are not inside parentheses."""
    out, depth, buf = [], 0, []
    for ch in text:
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth = max(0, depth - 1)
        if ch == "," and depth == 0:
            out.append("".join(buf))
            buf = []
        else:
            buf.append(ch)
    out.append("".join(buf))
    return [s.strip() for s in out if s.strip()]


def linkedin_skills(r: Resume) -> list[str]:
    """One LinkedIn-style tag per top-level skill, parenthetical detail dropped."""
    skills = r.section("Skills")
    flat: list[str] = []
    if skills:
        for it in skills.items:
            flat += [_PAREN.sub("", k).strip() for k in split_top_level(strip_inline(it.text))]
    return [k for k in flat if k]


def render_linkedin(r: Resume) -> str:
    """Paste-ready blocks. Limits: headline 220, about 2600, experience description 2000."""
    o = ["# LinkedIn paste blocks", "", "Copy each block into the matching LinkedIn field. "
         "Character counts are shown against LinkedIn's limits.", ""]

    def block(name: str, text: str, limit: int) -> None:
        flag = "OK" if len(text) <= limit else "TOO LONG"
        o.extend([f"## {name} ({len(text)}/{limit} chars, {flag})", "", "```", text, "```", ""])

    block("Headline", linkedin_headline(r), LINKEDIN_LIMITS["headline"])
    if r.section("Summary"):
        block("About", linkedin_about(r), LINKEDIN_LIMITS["about"])
    exp = r.section("Experience")
    if exp:
        for e in exp.entries:
            text, dropped = linkedin_entry_text(e)
            block(f"Experience — {e.title} ({e.org or 'no org'})", text, LINKEDIN_LIMITS["description"])
            if dropped:
                o.append(f"_{dropped} bullet(s) omitted to stay under LinkedIn's 2000-character limit; "
                         f"see the full list in the PDF or txt output._")
                o.append("")
            o.append(f"Fields: Title = `{e.title}` · Company = `{e.org}` · Dates = `{e.dates}` · Location = `{e.location}`")
            o.append("")
    proj = r.section("Projects")
    if proj:
        for e in proj.entries:
            text, _ = linkedin_entry_text(e, limit=None)
            block(f"Project — {e.title}", text, LINKEDIN_LIMITS["description"])
    if r.section("Skills"):
        o.extend(["## Skills (add individually; LinkedIn allows 100)", "", "```",
                  "\n".join(linkedin_skills(r)), "```", ""])
    return "\n".join(o)


# --------------------------------------------------------------------------- #
# Renderers: binary
# --------------------------------------------------------------------------- #

def render_pdf_and_jpg(html_text: str, pdf_path: Path | None, jpg_path: Path | None, r: Resume) -> None:
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = None
        for channel in ("msedge", "chrome", None):
            try:
                browser = p.chromium.launch(channel=channel) if channel else p.chromium.launch()
                break
            except Exception:
                continue
        if browser is None:
            sys.exit("No Chromium-based browser available for Playwright (tried Edge, Chrome, bundled).")
        page = browser.new_page(viewport={"width": 816, "height": 1056}, device_scale_factor=2)
        page.set_content(html_text, wait_until="load")
        page.emulate_media(media="print")
        if pdf_path:
            page.pdf(path=str(pdf_path), format="Letter", print_background=True, prefer_css_page_size=True)
        if jpg_path:
            page.emulate_media(media="screen")
            png = page.screenshot(full_page=True, type="png")
            from PIL import Image
            im = Image.open(io.BytesIO(png)).convert("RGB")
            im.save(jpg_path, "JPEG", quality=90, optimize=True)
        browser.close()

    if pdf_path:
        try:
            from pypdf import PdfReader, PdfWriter
            reader = PdfReader(str(pdf_path))
            writer = PdfWriter()
            writer.append(reader)
            skills = r.section("Skills")
            kw = ", ".join(strip_inline(i.text) for i in skills.items) if skills else ""
            writer.add_metadata({
                "/Title": f"{r.name} — Résumé ({r.meta.get('variant', 'general')})",
                "/Author": r.name,
                "/Subject": r.meta.get("title", ""),
                "/Keywords": kw[:1000],
                "/Creator": "resume/build.py",
            })
            with open(pdf_path, "wb") as fh:
                writer.write(fh)
        except ImportError:
            pass


def render_docx(r: Resume, path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for side in ("left_margin", "right_margin"):
        setattr(sec, side, Inches(0.7))
    sec.top_margin = sec.bottom_margin = Inches(0.6)
    text_width = Inches(8.5 - 1.4)

    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10.5)
    base.paragraph_format.space_after = Pt(2)
    doc.core_properties.author = r.name
    doc.core_properties.title = f"{r.name} — Résumé"
    doc.core_properties.subject = r.meta.get("title", "")
    doc.core_properties.keywords = r.meta.get("variant", "")

    def para(text_runs, *, size=None, bold=False, color=None, align=None, after=None, before=None, style=None):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        if isinstance(text_runs, str):
            text_runs = [(text_runs, bold)]
        for t, b in text_runs:
            run = p.add_run(t)
            run.bold = b or bold
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        if align:
            p.alignment = align
        if after is not None:
            p.paragraph_format.space_after = Pt(after)
        if before is not None:
            p.paragraph_format.space_before = Pt(before)
        return p

    para(r.name, size=20, bold=True, after=0)
    tl = " · ".join(x for x in [r.meta.get("title"), r.meta.get("location")] if x)
    if tl:
        para(tl, size=11, color="444444", after=0)
    para(" | ".join(v for v, _ in contact_items(r)), size=9.5, color="444444", after=6)

    for s in r.sections:
        h = para(s.title.upper(), size=11, bold=True, color="1F3A5F", before=8, after=2)
        # bottom border as a rule under each section heading
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = h._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "1F3A5F")):
            bottom.set(qn(k), v)
        pbdr.append(bottom)
        pPr.append(pbdr)

        if s.kind == "text":
            for ptxt in s.paragraphs:
                para(inline_runs(ptxt))
        elif s.kind == "list":
            for it in s.items:
                runs = ([(f"{it.label}: ", True)] if it.label else []) + inline_runs(it.text)
                para(runs, after=1)
        else:
            for e in s.entries:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
                run = p.add_run(e.title)
                run.bold = True
                if e.org:
                    p.add_run(f" — {e.org}")
                if e.dates:
                    d = p.add_run(f"\t{e.dates}")
                    d.font.color.rgb = RGBColor.from_string("444444")
                if e.location:
                    para(e.location, size=9.5, color="666666", after=0)
                for d in e.description:
                    para(inline_runs(d), after=1)
                for b in e.bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = Pt(1)
                    for t, bold in inline_runs(b):
                        rr = bp.add_run(t)
                        rr.bold = bold
    doc.save(str(path))


# --------------------------------------------------------------------------- #
# Driver
# --------------------------------------------------------------------------- #

def build_variant(src: Path, formats: list[str], terms: list[str], css: str,
                  theme: str | None = None) -> dict[str, Path]:
    text = src.read_text(encoding="utf-8")
    guard(src.name, text, terms)
    r = parse(text, src.name)
    variant = r.meta.get("variant") or src.stem
    stem = f"Steven-Shine-Resume-{variant}" + (f"-{theme}" if theme else "")
    outdir = OUT / variant
    outdir.mkdir(parents=True, exist_ok=True)

    text_outputs: dict[str, str] = {}
    if "md" in formats:
        text_outputs["md"] = text
    if "txt" in formats:
        text_outputs["txt"] = render_txt(r)
    if "json" in formats:
        text_outputs["json"] = render_json(r)
    html_text = render_html(r, css)
    if "html" in formats:
        text_outputs["html"] = html_text
    if "linkedin" in formats:
        text_outputs["linkedin"] = render_linkedin(r)

    # Guard every text rendering before touching disk (binary outputs derive from the same model).
    for ext, body in text_outputs.items():
        guard(f"{stem}.{ext}", body, terms)
    guard(f"{stem}.html(internal)", html_text, terms)

    written: dict[str, Path] = {}
    for ext, body in text_outputs.items():
        p = outdir / (f"{stem}-linkedin.md" if ext == "linkedin" else f"{stem}.{ext}")
        p.write_text(body, encoding="utf-8", newline="\n")
        written[ext] = p

    pdf_path = outdir / f"{stem}.pdf" if "pdf" in formats else None
    jpg_path = outdir / f"{stem}.jpg" if "jpg" in formats else None
    if pdf_path or jpg_path:
        render_pdf_and_jpg(html_text, pdf_path, jpg_path, r)
        if pdf_path:
            written["pdf"] = pdf_path
        if jpg_path:
            written["jpg"] = jpg_path
    if "docx" in formats:
        p = outdir / f"{stem}.docx"
        render_docx(r, p)
        written["docx"] = p
    return written


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("variants", nargs="*", help="variant names (src/<name>.md); default all")
    ap.add_argument("--formats", default=",".join(ALL_FORMATS), help="comma list of: " + ",".join(ALL_FORMATS))
    ap.add_argument("--check", action="store_true", help="denylist scan of sources only")
    ap.add_argument("--publish", metavar="VARIANT", help="copy VARIANT's PDF to dist/%s" % PUBLISHED_NAME)
    ap.add_argument("--theme", metavar="NAME", help="use themes/NAME.css instead of theme.css; "
                    "output names get a -NAME suffix")
    ap.add_argument("--list", action="store_true", help="list available variants")
    a = ap.parse_args(argv)

    sources = sorted(SRC.glob("*.md"))
    if a.list:
        for s in sources:
            print(s.stem)
        return 0
    if a.variants:
        wanted = set(a.variants)
        sources = [s for s in sources if s.stem in wanted]
        missing = wanted - {s.stem for s in sources}
        if missing:
            sys.exit(f"unknown variant(s): {', '.join(sorted(missing))}")
    if not sources:
        sys.exit(f"no sources found in {SRC}")

    terms = load_denylist()
    if a.check:
        for s in sources:
            guard(s.name, s.read_text(encoding="utf-8"), terms)
            parse(s.read_text(encoding="utf-8"), s.name)
            print(f"clean: {s.name}")
        return 0

    formats = [f.strip() for f in a.formats.split(",") if f.strip()]
    bad = [f for f in formats if f not in ALL_FORMATS]
    if bad:
        sys.exit(f"unknown format(s): {', '.join(bad)}")
    theme_file = (HERE / "themes" / f"{a.theme}.css") if a.theme else THEME
    if not theme_file.exists():
        sys.exit(f"theme not found: {theme_file}")
    css = theme_file.read_text(encoding="utf-8")

    for s in sources:
        written = build_variant(s, formats, terms, css, a.theme)
        print(f"{s.stem}:")
        for ext, p in written.items():
            print(f"  {ext:8s} {p.relative_to(HERE)}")

    if a.publish:
        suffix = f"-{a.theme}" if a.theme else ""
        src_pdf = OUT / a.publish / f"Steven-Shine-Resume-{a.publish}{suffix}.pdf"
        if not src_pdf.exists():
            sys.exit(f"--publish: {src_pdf} not built (include pdf in --formats and the variant in the list)")
        DIST.mkdir(exist_ok=True)
        shutil.copyfile(src_pdf, DIST / PUBLISHED_NAME)
        print(f"published {a.publish} -> {(DIST / PUBLISHED_NAME).relative_to(HERE)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
